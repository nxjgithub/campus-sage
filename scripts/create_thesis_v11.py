from __future__ import annotations

import json
import shutil
from pathlib import Path

from docx import Document


ROOT_DIR = Path(__file__).resolve().parents[1]
SOURCE_DOCX = ROOT_DIR / "论文" / "版本" / "CampusSage-毕业设计论文-V10.docx"
TARGET_DOCX = ROOT_DIR / "论文" / "版本" / "CampusSage-毕业设计论文-V11.docx"
SUMMARY_PATH = ROOT_DIR / "outputs" / "eval_official_formal_v11_fast" / "official_v11_fast_eval_summary.json"
MANIFEST_PATH = ROOT_DIR / "docs" / "examples" / "official_formal_corpus_v11_manifest.json"


def main() -> None:
    """生成对齐 V11 正式文件实验的论文版本。"""

    shutil.copyfile(SOURCE_DOCX, TARGET_DOCX)
    summary = json.loads(SUMMARY_PATH.read_text(encoding="utf-8"))
    manifest = json.loads(MANIFEST_PATH.read_text(encoding="utf-8"))
    doc = Document(TARGET_DOCX)
    paragraphs = doc.paragraphs

    def set_p(index: int, text: str) -> None:
        paragraphs[index].text = text

    # 摘要与英文摘要。
    set_p(44, "为验证系统有效性，本文进一步使用 data/正式文件 目录中的真实校园文件构建 V11 扩展实验集。系统共解析 45 份 PDF/DOCX 正式文件，生成 405 个文本块，并构造 100 条评测问题，其中 90 条为知识库内问题、10 条为知识库外边界问题。实验采用 BAAI/bge-small-zh-v1.5 生成文本向量，并使用 BAAI/bge-reranker-base CrossEncoder 对召回候选进行二阶段重排。在 TopK=5 条件下，系统 Recall@5 达到 0.9556，MRR 达到 0.9411；关闭重排时 Recall@5 为 0.9222、MRR 为 0.8381，说明重排模型能够明显提升证据排序质量。实验结果表明，CampusSage 能够在更接近真实校园资料规模的数据集上完成可追溯的证据检索、重排和问答支撑。")
    set_p(49, "This thesis designs and implements CampusSage, an evidence-grounded university knowledge assistant based on Retrieval-Augmented Generation (RAG). The system targets common campus information service scenarios such as academic affairs notices, graduation procedures, safety reminders, employment services, system manuals, student management, and campus traffic regulations.")
    set_p(50, "The backend is implemented with FastAPI and uses Qdrant as the default vector database. The retrieval pipeline supports OpenAI-compatible embedding services and local sentence-transformers models. The current V11 experiment uses BAAI/bge-small-zh-v1.5 for embeddings and BAAI/bge-reranker-base as a CrossEncoder reranker. The frontend is implemented with React and provides knowledge-base management, document ingestion, question answering, citation display, and monitoring-oriented interaction views.")
    set_p(51, "In the latest verified V11 experiment, the evaluation corpus contains 45 official campus files, 405 text chunks, and 100 evaluation questions, including 90 in-scope questions and 10 out-of-scope boundary questions. With TopK=5 and reranking enabled, the retrieval pipeline achieves Recall@5 = 0.9556 and MRR = 0.9411. Compared with retrieval without reranking, reranking improves Recall@5 from 0.9222 to 0.9556 and MRR from 0.8381 to 0.9411.")

    # 第三、五、六章中与实验设置和重排相关的表述。
    set_p(263, "本系统当前采用可配置重排模块，既支持 HTTP /rerank 服务，也支持本地 CrossEncoder 模型。V11 正式文件实验使用 BAAI/bge-reranker-base 对 query-document pair 进行相关性评分，从而对向量召回候选进行二阶段排序。")
    set_p(264, "将 Reranker 设计为可选模块，是出于工程实现和实验分析两方面的考虑。在小规模知识库中，仅使用向量 TopK 已能完成基本问答链路；当正式文件数量增加、相似通知增多时，开启 cross-encoder 重排可以把正确证据提前。V11 实验中，重排将 Recall@5 从 0.9222 提升到 0.9556，将 MRR 从 0.8381 提升到 0.9411。")
    set_p(373, "系统可行性还需要由实验结果支撑。V11 实验使用 45 份正式校园文件和 100 条评测问题，对 TopK、阈值和重排策略进行了比较。TopK=5 且开启 BAAI/bge-reranker-base 重排时，Recall@5 达到 0.9556，MRR 为 0.9411；TopK=3 时 Recall 为 0.9444，说明扩大到 TopK=5 能覆盖少量更靠后的标准证据。")
    set_p(489, "后端服务使用 FastAPI 实现，运行在本地 Python 虚拟环境与 Docker Compose 编排环境中。向量数据库默认使用 Qdrant，关系型数据使用 MySQL，任务队列与缓存使用 Redis，对象存储使用 MinIO。")
    set_p(490, "V11 实验在本地 CPU 环境下使用 sentence-transformers 加载 BAAI/bge-small-zh-v1.5 生成向量，并使用 BAAI/bge-reranker-base CrossEncoder 执行重排。系统工程配置仍支持通过 HTTP 服务接入 TEI embedding 与 /rerank 重排服务。")
    set_p(491, "实验默认采用 chunk_size=500、chunk_overlap=100、candidate_topk=20、TopK=5 的配置。检索流程先召回候选片段，再由 cross-encoder 重排模型计算相关性分数，最后统计标准文档是否进入前 K 个候选。")
    set_p(495, "本文使用 data/正式文件 目录中的真实校园文件构建 V11 扩展实验集。解析器共成功解析 45 份 PDF/DOCX 文件，覆盖教务、就业、征兵、安全、心理咨询、学生社区、交通管理、系统平台使用、学生证与公共服务等场景。评测集包含 100 条问题，其中 90 条为知识库内问题，10 条为知识库外边界问题。")
    set_p(497, "评测样本包含文件概述题、主题细节题和知识库外问题。每份正式文件至少对应 2 条知识库内问题，用于验证系统是否能够在更多相似校园通知中定位标准文档；知识库外问题用于验证阈值与拒答边界。")
    set_p(499, "本文主要使用 Recall@K、MRR、Top1 命中数、阈值拒答率和检索重排耗时作为核心评测指标。对于问答链路，还关注引用是否能够定位到 doc_name、page 或 section_path、snippet 等证据字段。")
    set_p(511, "从表 5-4 可以看出，当 TopK 从 3 增加到 5 时，Recall 从 0.9444 提升到 0.9556，说明部分标准证据位于第 4 至第 5 个候选之间。当 TopK 从 5 增加到 8 时，Recall 与 MRR 不再提升，因此 V11 默认采用 TopK=5。")
    set_p(514, "从表 5-5 可以看出，在 90 条知识库内问题中，TopK=5 能命中 86 条标准文档，Top1 命中 84 条。该结果说明在正式文件扩展后，系统仍能把大多数标准证据排在首位或前 5 位。")
    set_p(518, "文本分块会影响向量语义表达、证据片段完整性和上下文预算。V11 正式文件实验采用 chunk_size=500、chunk_overlap=100 的配置，45 份正式文件共生成 405 个文本块。")
    set_p(520, "表 5-6 展示了当前已验证配置以及后续可用于消融实验的候选配置。由于 V11 本轮重点验证正式文件规模扩展和重排有效性，其他分块参数未重建索引，不作为量化结论。")
    set_p(526, "阈值实验显示，在 threshold 为 none 与 0.25 时，知识库内 Recall@5 均为 0.9556，但知识库外问题不会被拒答；当 threshold 提高到 0.50 时，知识库内 Recall@5 仍保持 0.9556，同时知识库外拒答率提升到 0.5000。说明阈值可增强边界控制，但仍需要结合关键词覆盖率与最小上下文长度共同判断。")
    set_p(530, "重排用于对向量召回结果进行二阶段相关性排序。V11 实验使用 BAAI/bge-reranker-base CrossEncoder，对 query 与候选文本对进行相关性评分。")
    set_p(532, "表 5-8 展示了关闭与开启重排时的对比结果。关闭重排时 Recall@5 为 0.9222，MRR 为 0.8381；开启重排后 Recall@5 提升到 0.9556，MRR 提升到 0.9411。")
    set_p(533, "该结果说明，在 45 份正式文件、405 个文本块的扩展语料上，重排不只是工程预留能力，而是能够实际改善证据排序质量。尤其是 Top1 命中数从 72/90 提升到 84/90，更有利于后续答案生成阶段优先引用正确证据。")
    set_p(540, "从表 5-9 可以看出，直接大模型问答缺少知识库证据约束，无法稳定给出引用来源，也难以判断学校制度是否已经更新。CampusSage 虽然引入了检索和重排开销，但在 V11 正式文件实验中 Recall@5 达到 0.9556，能够为后续回答提供更可靠的证据基础。")
    set_p(549, "实验结果表明，系统在 V11 正式文件评测集上的 Recall@5 为 0.9556，MRR 为 0.9411，说明多数问题能够在前 5 个候选中召回标准证据，并且命中样本排序位置较靠前。")
    set_p(551, "V11 实验表明，BAAI/bge-reranker-base 对正式校园文件具有明显排序收益。相比仅使用向量相似度，重排使 Recall@5 从 0.9222 提升到 0.9556，使 MRR 从 0.8381 提升到 0.9411。")
    set_p(555, "本章围绕 CampusSage 的真实运行环境、正式文件语料、评测问题、检索参数、重排模型和结果分析展开实验验证。V11 实验使用 45 份正式校园文件和 100 条评测问题，在 TopK=5 且开启 BAAI/bge-reranker-base 重排时，Recall@5 达到 0.9556，MRR 达到 0.9411。实验说明系统能够在更接近真实校园资料规模的数据集上完成基于证据的校园知识问答核心流程。")
    set_p(562, "（4）完成了真实数据扩展评测与结果分析。系统在 45 份正式校园文件和 100 条评测问题上完成验证，TopK=5 且开启 BAAI/bge-reranker-base 重排时，Recall@5 为 0.9556，MRR 为 0.9411。")
    set_p(572, "首先，当前语料规模相比早期演示集已经明显扩大。第五章 V11 实验使用 45 份正式校园文件和 100 条评测问题，能够更充分地说明系统链路是有效的，也能看出 TopK、阈值和重排策略对结果的影响。真实环境中仍会有更多学院通知、历史制度、临时公告、附件表格和不同年份版本，因此后续还需要继续扩充语料并建立版本治理机制。")
    set_p(575, "（2）重排收益已经在正式文件扩展实验中得到验证，但仍需继续优化推理成本。当前系统已支持 BAAI/bge-reranker-base CrossEncoder 重排，在 V11 数据集上能明显提升 Recall 与 MRR。后续需要结合批量大小、缓存、候选池规模和服务部署方式降低延迟。")
    set_p(580, "第一，可以继续从扩充语料和版本治理做起。现阶段 45 份正式文件和 100 条问题已经能验证系统原型，但真实校园场景还应逐步纳入更多学院制度、历史公告、附件表格和不同年份版本。")
    set_p(583, "（3）检索与重排优化。后续可以从候选池规模、批量大小、重排缓存、模型服务化部署和难例数据集构造等方面继续优化，使重排在复杂问题和大规模知识库中发挥更稳定作用。")

    update_tables(doc, summary, manifest)
    doc.save(TARGET_DOCX)
    print(f"created: {TARGET_DOCX}")


def set_cell(table, row: int, column: int, text: str) -> None:
    table.rows[row].cells[column].text = text


def trim_rows(table, count: int) -> None:
    while len(table.rows) > count:
        table._tbl.remove(table.rows[-1]._tr)


def ensure_rows(table, count: int) -> None:
    while len(table.rows) < count:
        table.add_row()


def update_rows(table, rows: list[tuple[str, ...]]) -> None:
    trim_rows(table, len(rows))
    ensure_rows(table, len(rows))
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            set_cell(table, row_index, column_index, value)


def update_tables(doc: Document, summary: dict, manifest: dict) -> None:
    """更新第五章实验表格。"""

    set_cell(doc.tables[4], 8, 0, "默认检索与重排参数")
    set_cell(doc.tables[4], 8, 1, "candidate_topk=20，TopK=5，rerank_enabled=true")
    set_cell(doc.tables[4], 8, 2, "向量召回后通过 BAAI/bge-reranker-base CrossEncoder 重排")
    set_cell(doc.tables[4], 9, 1, "chunk_size=500，chunk_overlap=100")
    set_cell(doc.tables[4], 9, 2, "V11 正式文件实验共生成 405 个文本块")
    set_cell(doc.tables[4], 10, 1, "45 份正式文件，100 条问题")
    set_cell(doc.tables[4], 10, 2, "其中 90 条知识库内问题，10 条知识库外边界问题")

    corpus_rows = [("序号", "文档名称", "字符数", "chunk 数")]
    for index, item in enumerate(manifest["documents"], start=1):
        corpus_rows.append(
            (
                str(index),
                item["doc_name"],
                str(item["char_count"]),
                str(item["chunk_count"]),
            )
        )
    update_rows(doc.tables[5], corpus_rows)

    topk = summary["topk_results"]
    update_rows(
        doc.tables[6],
        [
            ("组别", "TopK", "Recall@K", "MRR", "命中数/样本"),
            ("A", "3", str(topk[0]["recall"]), str(topk[0]["mrr"]), f"{topk[0]['hit_count']}/{topk[0]['samples']}"),
            ("B", "5", str(topk[1]["recall"]), str(topk[1]["mrr"]), f"{topk[1]['hit_count']}/{topk[1]['samples']}"),
            ("C", "8", str(topk[2]["recall"]), str(topk[2]["mrr"]), f"{topk[2]['hit_count']}/{topk[2]['samples']}"),
        ],
    )
    update_rows(
        doc.tables[7],
        [
            ("TopK", "candidate_topk", "final_hit_count", "Top1命中数", "平均耗时/ms", "P95耗时/ms"),
            ("3", "20", "85/90", "84/90", "3246", "4094"),
            ("5", "20", "86/90", "84/90", "3246", "4094"),
            ("8", "20", "86/90", "84/90", "3246", "4094"),
        ],
    )
    update_rows(
        doc.tables[8],
        [
            ("组别", "chunk_size", "chunk_overlap", "chunk 数", "Recall@5", "MRR", "平均检索耗时/ms", "备注"),
            ("当前默认", "500", "100", "405", "0.9556", "0.9411", "3246", "V11 正式文件实验配置"),
            ("细粒度候选", "300", "50", "未重建索引", "-", "-", "-", "后续消融实验配置"),
            ("大粒度候选", "800", "150", "未重建索引", "-", "-", "-", "后续消融实验配置"),
        ],
    )
    threshold = summary["threshold_results"]
    update_rows(
        doc.tables[9],
        [
            ("组别", "threshold", "Recall@5", "知识库内拒答率", "知识库外拒答率", "accepted_count", "samples"),
            ("A", "none", str(threshold[0]["in_scope_recall_at_5"]), "0.0", "0.0", "100", "100"),
            ("B", "0.25", str(threshold[1]["in_scope_recall_at_5"]), "0.0", "0.0", "100", "100"),
            ("C", "0.50", str(threshold[2]["in_scope_recall_at_5"]), "0.0", "0.5", "95", "100"),
        ],
    )
    rerank = summary["rerank_compare"]
    update_rows(
        doc.tables[10],
        [
            ("组别", "是否重排", "Recall@5", "MRR", "Top1命中数", "命中数/样本"),
            ("A", "否", str(rerank[0]["recall"]), str(rerank[0]["mrr"]), "72/90", "83/90"),
            ("B", "是", str(rerank[1]["recall"]), str(rerank[1]["mrr"]), "84/90", "86/90"),
        ],
    )
    update_rows(
        doc.tables[11],
        [
            ("方法", "是否依赖知识库证据", "是否返回引用", "是否支持知识库外拒答", "本轮量化指标", "说明"),
            ("直接大模型问答", "否", "否", "否", "未单独量化", "无法证明答案来自校园知识库最新制度"),
            ("CampusSage RAG", "是", "是", "是", "Recall@5=0.9556，MRR=0.9411", "通过向量召回、CrossEncoder 重排、证据过滤与引用展示约束回答"),
        ],
    )


if __name__ == "__main__":
    main()
